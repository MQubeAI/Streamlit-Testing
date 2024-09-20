# from docx import Document

# def replace_placeholder_in_docx(template_path, output_path, player_name):
#     # Load the template .docx file
#     doc = Document(template_path)
    
#     # Iterate through each paragraph in the document
#     for paragraph in doc.paragraphs:
#         # If the placeholder exists in the paragraph, replace it
#         if '{Name_of_Player}' in paragraph.text:
#             paragraph.text = paragraph.text.replace('{Name_of_Player}', player_name)
    
#     # Save the new document with the replaced placeholder
#     doc.save(output_path)
#     print(f"New document generated: {output_path}")

# # Define paths and player name
# template_file = 'template.docx'  # Path to your template
# output_file = 'placeholder_contract.docx'  # Path for the new generated file
# player_name = 'Lionel Messi'

# replace_placeholder_in_docx(template_file, output_file, player_name)


#inputs





from docx import Document

def replace_placeholder_in_docx(template_path, output_path, values_dict):

    doc = Document(template_path)
    

    for paragraph in doc.paragraphs:
        
        for placeholder, replacement in values_dict.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, replacement)
    
    doc.save(output_path)
    print(f"New document generated: {output_path}")


template_file = 'template1.docx'  
output_file = 'player_contract.docx'  

values_dict = {
    '{player_name}': 'Serdar Yildiz',
    '{player_nationality}': 'Turkish',
    '{player_address}': '14 A lal kuan Noida',
    '{player_passport_no}': 'CLUPC7537H',
    '{number_of_seasons}': 'Total 3 Seasons, First, Second, Third Season',
    '{signature_date}': '31.05.24',
    '{end_date_of_contract}': '1.06.2026',
    '{First_season_salary}': '250000 EURO',
    '{Second_season_salary}': '20000 EURO',
    '{Third_season_salary}': '150000 EURO',
    '{Attendance_fee_for_first_Season}': '20000 EURO',
    '{Attendance_fee_for_second_Season}': '10000 EURO',
    '{Attendance_fee_for_third_Season}': '150000 EURO',
    '{Bonuses_amount_for_first_Season}': '20000 EURO',
    '{Bonuses_amount_for_second_Season}': '20000 EURO',
    '{Bonuses_amount_for_third_Season}': '20000 EURO',
    '{Signing_Fee}': '20000 EURO'
}

replace_placeholder_in_docx(template_file, output_file, values_dict)



